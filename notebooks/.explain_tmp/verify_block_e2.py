import os, tempfile

print("=== ch26 cell5 variant: fpdf with '-' instead of em dash ===")
from fpdf import FPDF
import pdfplumber
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
pdf.cell(200, 10, text="Srivatsa Gorti", new_x="LMARGIN", new_y="NEXT", align="C")
pdf.set_font("Arial", size=10)
pdf.cell(200, 10, text="srivatsa@email.com", new_x="LMARGIN", new_y="NEXT", align="C")
pdf.set_font("Arial", style="B", size=11)
pdf.cell(200, 10, text="PROFESSIONAL SUMMARY", new_x="LMARGIN", new_y="NEXT")
pdf.set_font("Arial", size=10)
pdf.multi_cell(0, 5, text="Data scientist with 5+ years of Python, NLP, and ML experience.")
pdf.set_font("Arial", style="B", size=11)
pdf.cell(200, 10, text="EXPERIENCE", new_x="LMARGIN", new_y="NEXT")
pdf.set_font("Arial", size=10)
pdf.cell(200, 10, text="Google - Senior Data Scientist, 2020-Present", new_x="LMARGIN", new_y="NEXT")
test_pdf = os.path.join(tempfile.gettempdir(), "test_resume.pdf")
pdf.output(test_pdf)
print(f"PDF created: {os.path.getsize(test_pdf)} bytes")
with pdfplumber.open(test_pdf) as pdfp:
    for i, page in enumerate(pdfp.pages):
        text = page.extract_text() or ""
        print(f"\nPage {i+1}:\n{text[:200]}")

print("\n=== ch27 cell4: docx round-trip ===")
from docx import Document
from docx.shared import Pt
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
doc.add_heading('Srivatsa Gorti', level=0)
doc.add_paragraph('srivatsa@email.com | +1-555-1234')
doc.add_heading('PROFESSIONAL SUMMARY', level=1)
doc.add_paragraph('Data scientist with 5+ years experience in Python, NLP, and ML.')
doc.add_heading('EXPERIENCE', level=1)
p = doc.add_paragraph()
p.add_run('Google, Mountain View').bold = True
p.add_run(' - Senior Data Scientist')
doc.add_paragraph('Jan 2020 - Present')
doc.add_paragraph('Developed NLP pipelines processing 10M+ documents daily')
test_docx = os.path.join(tempfile.gettempdir(), "test_resume.docx")
doc.save(test_docx)
print(f"DOCX bytes: {os.path.getsize(test_docx)}")
doc2 = Document(test_docx)
print("=== Extracted Content ===")
for para in doc2.paragraphs:
    if para.text.strip():
        print(f"  [{para.style.name:20s}] {para.text[:60]}")

print("\n=== ch27 cell6: table extraction ===")
doc3 = Document()
doc3.add_heading('Skills Matrix', level=1)
table = doc3.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
for i, (skill, level_, years) in enumerate([
    ("Python", "Expert", "5"), ("TensorFlow", "Advanced", "3"), ("NLP", "Advanced", "4"),
]):
    table.rows[i+1].cells[0].text = skill
    table.rows[i+1].cells[1].text = level_
    table.rows[i+1].cells[2].text = years
table.rows[0].cells[0].text = "Skill"
table.rows[0].cells[1].text = "Level"
table.rows[0].cells[2].text = "Years"
doc3.save(os.path.join(tempfile.gettempdir(), "skills.docx"))
doc4 = Document(os.path.join(tempfile.gettempdir(), "skills.docx"))
for i, table in enumerate(doc4.tables):
    print(f"Table {i+1}: {len(table.rows)} rows x {len(table.columns)} cols")
    for row in table.rows:
        print(f"  {[cell.text for cell in row.cells]}")

print("\n=== ch28 cell6: preprocessing ===")
from PIL import Image, ImageEnhance
try:
    img = Image.open(os.path.join(tempfile.gettempdir(), "test_resume.png"))
    gray = img.convert("L")
    enhancer = ImageEnhance.Contrast(gray)
    high_contrast = enhancer.enhance(2.0)
    binary = gray.point(lambda x: 255 if x > 128 else 0)
    print(f"img mode={img.mode} size={img.size}; gray mode={gray.mode}; binary mode={binary.mode}")
    print("Preprocessing: grayscale + contrast + threshold = better OCR results")
except Exception as e:
    print(f"ch28 cell6 FAILED: {type(e).__name__}: {e}")
